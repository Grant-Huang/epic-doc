"""EpicDoc — the main fluent document builder."""
from __future__ import annotations

import io
from typing import Any, Dict, List, Optional, Union

from docx import Document
from docx.shared import Pt

from epic_doc.styles import get_theme
from epic_doc.styles.theme import Theme
from epic_doc.utils.tempfiles import TempFileManager


class EpicDoc:
    """Fluent builder for complex, beautifully themed DOCX documents.

    Usage::

        from epic_doc import EpicDoc

        doc = EpicDoc(theme='professional')
        doc.add_heading('Annual Report', level=1)
        doc.add_paragraph('This report covers Q1–Q4 2026.')
        doc.add_table(data=[['Region', 'Revenue'], ['APAC', 1200]], headers=True)
        doc.add_chart(chart_type='bar', data={'Q1': 100, 'Q2': 200}, title='Revenue')
        doc.save('report.docx')

    All ``add_*`` methods return ``self``, enabling method chaining.
    """

    def __init__(
        self,
        theme: Union[str, Theme] = "professional",
        template: Optional[str] = None,
    ) -> None:
        """Create a new EpicDoc instance.

        Args:
            theme: Theme name string (e.g. ``'professional'``) or a :class:`Theme` instance.
            template: Path to an existing ``.docx`` file to use as a style template.
                      When supplied, the document inherits all styles from that file.
        """
        if isinstance(theme, str):
            self._theme: Theme = get_theme(theme)
        else:
            self._theme = theme

        if template:
            self._doc = Document(template)
        else:
            self._doc = Document()
            self._apply_base_styles()

        self._tmgr = TempFileManager()

    # ── Base style setup ──────────────────────────────────────────────────────

    def _apply_base_styles(self) -> None:
        """Apply theme fonts and spacing to the document's built-in styles."""
        from docx.shared import RGBColor

        def _rgb(h: str) -> RGBColor:
            h = h.lstrip("#")
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

        t = self._theme
        styles = self._doc.styles

        # Normal style
        normal = styles["Normal"]
        normal.font.name = t.body_font
        normal.font.size = Pt(t.body_size)
        normal.font.color.rgb = _rgb(t.body_text)

        # Heading styles
        _heading_cfg = {
            "Heading 1": (t.h1_size, t.primary,   t.heading_font, t.h1_bold),
            "Heading 2": (t.h2_size, t.secondary, t.heading_font, t.h2_bold),
            "Heading 3": (t.h3_size, t.accent,    t.heading_font, t.h3_bold),
            "Heading 4": (t.h4_size, t.accent,    t.heading_font, t.h4_bold),
        }
        for style_name, (size, color, font, bold) in _heading_cfg.items():
            try:
                s = styles[style_name]
                s.font.name = font
                s.font.size = Pt(size)
                s.font.color.rgb = _rgb(color)
                s.font.bold = bold
            except KeyError:
                pass

        # Page margins (comfortable reading width)
        for section in self._doc.sections:
            from docx.shared import Inches
            section.left_margin   = Inches(1.18)
            section.right_margin  = Inches(1.18)
            section.top_margin    = Inches(1.0)
            section.bottom_margin = Inches(1.0)

    # ── Metadata ─────────────────────────────────────────────────────────────

    def set_metadata(
        self,
        title: Optional[str] = None,
        author: Optional[str] = None,
        subject: Optional[str] = None,
        description: Optional[str] = None,
    ) -> "EpicDoc":
        """Set document metadata (title, author, subject, description)."""
        from epic_doc.elements.layout import set_metadata
        set_metadata(self._doc, title=title, author=author,
                     subject=subject, description=description)
        return self

    # ── Text elements ─────────────────────────────────────────────────────────

    def add_heading(
        self,
        text: str,
        level: int = 1,
        align: str = "left",
    ) -> "EpicDoc":
        """Add a heading (level 1–4)."""
        from epic_doc.elements.text import add_heading
        add_heading(self._doc, self._theme, text, level=level, align=align)
        return self

    def add_paragraph(
        self,
        text: str,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        color: Optional[str] = None,
        align: str = "left",
        font_size: Optional[int] = None,
    ) -> "EpicDoc":
        """Add a body paragraph."""
        from epic_doc.elements.text import add_paragraph
        add_paragraph(
            self._doc, self._theme, text,
            bold=bold, italic=italic, underline=underline,
            color=color, align=align, font_size=font_size,
        )
        return self

    def add_list(
        self,
        items: List[Union[str, list]],
        style: str = "bullet",
    ) -> "EpicDoc":
        """Add a bullet or numbered list. Nested lists supported via nested Python lists."""
        from epic_doc.elements.text import add_list
        add_list(self._doc, self._theme, items, style=style)
        return self

    def add_hyperlink(self, text: str, url: str) -> "EpicDoc":
        """Add a paragraph containing a styled hyperlink."""
        from epic_doc.elements.text import add_hyperlink
        add_hyperlink(self._doc, self._theme, text, url)
        return self

    def add_code_block(self, code: str, language: Optional[str] = None) -> "EpicDoc":
        """Add a monospace code block with theme-colored background."""
        from epic_doc.elements.text import add_code_block
        add_code_block(self._doc, self._theme, code, language=language)
        return self

    def add_callout(
        self,
        text: str,
        style: str = "info",
        title: Optional[str] = None,
    ) -> "EpicDoc":
        """Add a highlighted callout box.

        style: info | warning | danger | success
        """
        from epic_doc.elements.text import add_callout
        add_callout(self._doc, self._theme, text, style=style, title=title)
        return self

    def add_horizontal_rule(self) -> "EpicDoc":
        """Add a thin horizontal rule line."""
        from epic_doc.elements.text import add_horizontal_rule
        add_horizontal_rule(self._doc, self._theme)
        return self

    # ── Table ─────────────────────────────────────────────────────────────────

    def add_table(
        self,
        data: List[List[Any]],
        headers: bool = True,
        style: str = "striped",
        col_widths: Optional[List[float]] = None,
        merge: Optional[List] = None,
        caption: Optional[str] = None,
        align: str = "left",
        font_size: Optional[int] = None,
        cell_align: str = "left",
    ) -> "EpicDoc":
        """Add a feature-rich styled table.

        Args:
            data: 2-D list (row-major). First row is header when ``headers=True``.
            headers: Style first row as dark header.
            style: striped | grid | minimal | bordered | dark
            col_widths: Column widths in inches.
            merge: List of (start_row, start_col, end_row, end_col) tuples.
            caption: Optional italic caption below the table.
            align: Table alignment — left | center | right.
            font_size: Override cell font size.
            cell_align: Default text alignment inside cells.
        """
        from epic_doc.elements.table import add_table
        add_table(
            self._doc, self._theme, data,
            headers=headers, style=style,
            col_widths=col_widths, merge=merge,
            caption=caption, align=align,
            font_size=font_size, cell_align=cell_align,
        )
        return self

    # ── Charts ────────────────────────────────────────────────────────────────

    def add_chart(
        self,
        chart_type: str = "bar",
        data: Optional[Dict] = None,
        title: Optional[str] = None,
        xlabel: Optional[str] = None,
        ylabel: Optional[str] = None,
        width: float = 5.5,
        height: float = 3.2,
        colors: Optional[List[str]] = None,
        caption: Optional[str] = None,
        show_values: bool = False,
        show_grid: bool = True,
        legend: bool = True,
    ) -> "EpicDoc":
        """Render a matplotlib chart and embed it.

        chart_type: bar | hbar | line | area | pie | scatter | combo
        data: Single-series ``{label: value}`` or multi-series ``{series: {label: value}}``.
        """
        from epic_doc.elements.chart import add_chart
        add_chart(
            self._doc, self._theme, self._tmgr,
            chart_type=chart_type, data=data or {},
            title=title, xlabel=xlabel, ylabel=ylabel,
            width=width, height=height,
            colors=colors, caption=caption,
            show_values=show_values, show_grid=show_grid,
            legend=legend,
        )
        return self

    # ── Diagrams ──────────────────────────────────────────────────────────────

    def add_flowchart(
        self,
        nodes: Dict,
        edges: List,
        direction: str = "TB",
        width: float = 5.0,
        caption: Optional[str] = None,
        graph_attrs: Optional[Dict] = None,
    ) -> "EpicDoc":
        """Render a graphviz flowchart and embed it.

        nodes: ``{node_id: label}`` or ``{node_id: {label, shape, color, style}}``.
        edges: List of ``(from, to)`` or ``(from, to, {label, style, color})``.
        direction: TB | LR | RL | BT
        """
        from epic_doc.elements.diagram import add_flowchart
        add_flowchart(
            self._doc, self._theme, self._tmgr,
            nodes=nodes, edges=edges,
            direction=direction, width=width,
            caption=caption, graph_attrs=graph_attrs,
        )
        return self

    # ── Images ────────────────────────────────────────────────────────────────

    def add_image(
        self,
        source: Union[str, bytes],
        width: Optional[float] = None,
        height: Optional[float] = None,
        align: str = "center",
        caption: Optional[str] = None,
    ) -> "EpicDoc":
        """Insert an image from a file path or raw bytes."""
        from epic_doc.elements.image import add_image
        add_image(self._doc, self._theme, source,
                  width=width, height=height, align=align, caption=caption)
        return self

    # ── Layout ────────────────────────────────────────────────────────────────

    def add_toc(self, title: str = "目录", depth: int = 3) -> "EpicDoc":
        """Insert a Table of Contents placeholder."""
        from epic_doc.elements.layout import add_toc
        add_toc(self._doc, self._theme, title=title, depth=depth)
        return self

    def set_header(
        self,
        text: str,
        align: str = "right",
    ) -> "EpicDoc":
        """Set the document header text."""
        from epic_doc.elements.layout import set_header
        set_header(self._doc, self._theme, text, align=align)
        return self

    def set_footer(
        self,
        text: Optional[str] = None,
        page_number: bool = True,
        align: str = "center",
    ) -> "EpicDoc":
        """Set the document footer, with optional page number."""
        from epic_doc.elements.layout import set_footer
        set_footer(self._doc, self._theme, text=text,
                   page_number=page_number, align=align)
        return self

    def add_page_break(self) -> "EpicDoc":
        """Insert an explicit page break."""
        from epic_doc.elements.layout import add_page_break
        add_page_break(self._doc)
        return self

    def add_section_break(self, break_type: str = "next_page") -> "EpicDoc":
        """Insert a section break (next_page | even_page | odd_page | continuous)."""
        from epic_doc.elements.layout import add_section_break
        add_section_break(self._doc, break_type=break_type)
        return self

    # ── Output ────────────────────────────────────────────────────────────────

    def save(self, path: str) -> None:
        """Save the document to a file path and clean up temp files."""
        try:
            self._doc.save(path)
        finally:
            self._tmgr.cleanup()

    def to_bytes(self) -> bytes:
        """Serialize the document to bytes (for HTTP responses, memory transfer, etc.)."""
        try:
            buf = io.BytesIO()
            self._doc.save(buf)
            return buf.getvalue()
        finally:
            self._tmgr.cleanup()

    # ── Context manager support ───────────────────────────────────────────────

    def __enter__(self) -> "EpicDoc":
        return self

    def __exit__(self, *_: object) -> None:
        self._tmgr.cleanup()

    # ── Repr ─────────────────────────────────────────────────────────────────

    def __repr__(self) -> str:
        return f"EpicDoc(theme='{self._theme.name}')"
