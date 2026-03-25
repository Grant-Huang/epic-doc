"""Chart element — renders matplotlib charts as PNG images embedded in DOCX."""
from __future__ import annotations

import io
from typing import TYPE_CHECKING, Any, Dict, List, Optional, Union

if TYPE_CHECKING:
    from docx.document import Document
    from epic_doc.styles.theme import Theme
    from epic_doc.utils.tempfiles import TempFileManager

# Single-series data: dict[label, value]
# Multi-series data:  dict[series_name, dict[label, value]]
DataType = Union[Dict[str, Any], Dict[str, Dict[str, Any]]]


def _configure_cjk_font() -> None:
    """Try to find and register a CJK-compatible font for matplotlib.

    Falls back silently if no CJK font is available — charts still render,
    but CJK characters may appear as boxes on font-limited systems.
    """
    try:
        import matplotlib.font_manager as fm

        # Preferred CJK fonts, in order
        _CJK_CANDIDATES = [
            "Noto Sans CJK SC", "Noto Sans SC", "Source Han Sans CN",
            "WenQuanYi Micro Hei", "Droid Sans Fallback",
            "PingFang SC", "Microsoft YaHei", "SimHei", "Arial Unicode MS",
        ]
        available = {f.name for f in fm.fontManager.ttflist}
        for name in _CJK_CANDIDATES:
            if name in available:
                import matplotlib as mpl
                mpl.rcParams["font.sans-serif"] = [name] + mpl.rcParams["font.sans-serif"]
                mpl.rcParams["axes.unicode_minus"] = False
                return
    except Exception:
        pass  # Never fail silently — chart rendering continues regardless


def _is_multiseries(data: DataType) -> bool:
    return data and isinstance(next(iter(data.values())), dict)


def _add_caption(doc: "Document", theme: "Theme", caption: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    def _rgb(h: str) -> RGBColor:
        h = h.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(caption)
    run.font.italic = True
    run.font.size = Pt(theme.caption_size)
    run.font.color.rgb = _rgb(theme.light_text)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)


def add_chart(
    doc: "Document",
    theme: "Theme",
    tmgr: "TempFileManager",
    chart_type: str = "bar",
    data: Optional[DataType] = None,
    title: Optional[str] = None,
    xlabel: Optional[str] = None,
    ylabel: Optional[str] = None,
    width: float = 5.5,
    height: float = 3.2,
    colors: Optional[List[str]] = None,
    caption: Optional[str] = None,
    show_values: bool = False,
    show_grid: bool = True,
    legend: bool = True,
) -> None:
    """Render a matplotlib chart and embed it in the document.

    Supported chart_type values:
        bar, hbar, line, pie, area, scatter, combo (bar + line overlay)
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        import matplotlib.ticker as mticker
    except ImportError as exc:
        raise ImportError(
            "matplotlib is required for add_chart(). "
            "Install it with: pip install matplotlib"
        ) from exc

    _configure_cjk_font()

    if data is None:
        data = {}

    palette = colors or theme.chart_palette
    fig, ax = plt.subplots(figsize=(width, height))
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")

    multiseries = _is_multiseries(data)

    # ── Single-series charts ─────────────────────────────────────────────────
    if chart_type == "bar":
        _render_bar(ax, data, palette, show_values, show_grid, multiseries)
    elif chart_type == "hbar":
        _render_hbar(ax, data, palette, show_values, show_grid, multiseries)
    elif chart_type == "line":
        _render_line(ax, data, palette, show_grid, multiseries)
    elif chart_type == "area":
        _render_area(ax, data, palette, show_grid, multiseries)
    elif chart_type == "pie":
        _render_pie(ax, data, palette)
    elif chart_type == "scatter":
        _render_scatter(ax, data, palette)
    elif chart_type == "combo":
        _render_combo(ax, data, palette, show_grid)
    else:
        raise ValueError(
            f"Unknown chart_type '{chart_type}'. "
            "Valid: bar, hbar, line, area, pie, scatter, combo"
        )

    # ── Decoration ───────────────────────────────────────────────────────────
    if title:
        ax.set_title(title, fontsize=11, fontweight="bold",
                     color="#" + theme.primary, pad=10)
    if xlabel:
        ax.set_xlabel(xlabel, fontsize=9, color="#" + theme.body_text)
    if ylabel:
        ax.set_ylabel(ylabel, fontsize=9, color="#" + theme.body_text)

    if legend and multiseries and chart_type != "pie":
        ax.legend(fontsize=8, framealpha=0.8)

    ax.tick_params(labelsize=8, colors="#" + theme.body_text)
    for spine in ax.spines.values():
        spine.set_edgecolor("#" + theme.table_grid)

    plt.tight_layout(pad=0.8)

    # ── Save & embed ─────────────────────────────────────────────────────────
    path = tmgr.new_png()
    fig.savefig(path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)

    from docx.shared import Inches
    doc.add_picture(path, width=Inches(width))

    if caption:
        _add_caption(doc, theme, caption)


# ── Render helpers ────────────────────────────────────────────────────────────

def _render_bar(ax, data, palette, show_values, show_grid, multiseries):
    import numpy as np
    if multiseries:
        series_names = list(data.keys())
        labels = list(next(iter(data.values())).keys())
        x = np.arange(len(labels))
        w = 0.8 / len(series_names)
        for i, (sname, sdata) in enumerate(data.items()):
            vals = [sdata.get(l, 0) for l in labels]
            bars = ax.bar(x + i * w - 0.4 + w / 2, vals, w,
                          label=sname, color=palette[i % len(palette)], alpha=0.85)
            if show_values:
                ax.bar_label(bars, fmt="%.1f", fontsize=7, padding=2)
        ax.set_xticks(x)
        ax.set_xticklabels(labels)
    else:
        labels = list(data.keys())
        vals   = [float(v) for v in data.values()]
        colors_  = [palette[i % len(palette)] for i in range(len(labels))]
        bars = ax.bar(labels, vals, color=colors_, alpha=0.85)
        if show_values:
            ax.bar_label(bars, fmt="%.1f", fontsize=7, padding=2)
    if show_grid:
        ax.yaxis.grid(True, linestyle="--", alpha=0.5)
        ax.set_axisbelow(True)


def _render_hbar(ax, data, palette, show_values, show_grid, multiseries):
    if multiseries:
        import numpy as np
        series_names = list(data.keys())
        labels = list(next(iter(data.values())).keys())
        y = list(range(len(labels)))
        h = 0.8 / len(series_names)
        for i, (sname, sdata) in enumerate(data.items()):
            vals = [sdata.get(l, 0) for l in labels]
            bars = ax.barh(
                [yi + i * h - 0.4 + h / 2 for yi in y],
                vals, h, label=sname,
                color=palette[i % len(palette)], alpha=0.85
            )
            if show_values:
                ax.bar_label(bars, fmt="%.1f", fontsize=7, padding=2)
        ax.set_yticks(y)
        ax.set_yticklabels(labels)
    else:
        labels = list(data.keys())
        vals   = [float(v) for v in data.values()]
        colors_ = [palette[i % len(palette)] for i in range(len(labels))]
        bars = ax.barh(labels, vals, color=colors_, alpha=0.85)
        if show_values:
            ax.bar_label(bars, fmt="%.1f", fontsize=7, padding=2)
    if show_grid:
        ax.xaxis.grid(True, linestyle="--", alpha=0.5)
        ax.set_axisbelow(True)


def _render_line(ax, data, palette, show_grid, multiseries):
    if multiseries:
        for i, (sname, sdata) in enumerate(data.items()):
            ax.plot(list(sdata.keys()), list(sdata.values()),
                    marker="o", linewidth=2, markersize=4,
                    color=palette[i % len(palette)], label=sname)
    else:
        ax.plot(list(data.keys()), [float(v) for v in data.values()],
                marker="o", linewidth=2, markersize=4, color=palette[0])
    if show_grid:
        ax.yaxis.grid(True, linestyle="--", alpha=0.5)
        ax.set_axisbelow(True)


def _render_area(ax, data, palette, show_grid, multiseries):
    if multiseries:
        for i, (sname, sdata) in enumerate(data.items()):
            c = palette[i % len(palette)]
            ax.fill_between(list(sdata.keys()), [float(v) for v in sdata.values()],
                            alpha=0.4, color=c, label=sname)
            ax.plot(list(sdata.keys()), [float(v) for v in sdata.values()],
                    linewidth=1.5, color=c)
    else:
        ax.fill_between(list(data.keys()), [float(v) for v in data.values()],
                        alpha=0.5, color=palette[0])
        ax.plot(list(data.keys()), [float(v) for v in data.values()],
                linewidth=1.5, color=palette[0])
    if show_grid:
        ax.yaxis.grid(True, linestyle="--", alpha=0.4)
        ax.set_axisbelow(True)


def _render_pie(ax, data, palette):
    labels = list(data.keys())
    vals   = [float(v) for v in data.values()]
    colors_ = [palette[i % len(palette)] for i in range(len(labels))]
    wedges, texts, autotexts = ax.pie(
        vals, labels=labels, colors=colors_,
        autopct="%1.1f%%", startangle=90,
        pctdistance=0.8, textprops={"fontsize": 8}
    )
    for at in autotexts:
        at.set_fontsize(7)
    ax.axis("equal")


def _render_scatter(ax, data, palette):
    """data expected as dict[series_name, list of (x, y)] or dict[label, value]."""
    if _is_multiseries(data):
        for i, (sname, points) in enumerate(data.items()):
            if isinstance(points, dict):
                xs = list(range(len(points)))
                ys = [float(v) for v in points.values()]
            else:
                xs, ys = zip(*points)
            ax.scatter(xs, ys, label=sname,
                       color=palette[i % len(palette)], alpha=0.8, s=40)
    else:
        xs = list(range(len(data)))
        ys = [float(v) for v in data.values()]
        ax.scatter(xs, ys, color=palette[0], alpha=0.8, s=40)


def _render_combo(ax, data, palette, show_grid):
    """Combo: first series as bar, rest as lines."""
    import numpy as np
    if not _is_multiseries(data):
        _render_bar(ax, data, palette, False, show_grid, False)
        return

    series = list(data.items())
    labels = list(series[0][1].keys())
    x = np.arange(len(labels))

    # First series → bar
    bar_vals = [float(series[0][1].get(l, 0)) for l in labels]
    ax.bar(x, bar_vals, color=palette[0], alpha=0.7, label=series[0][0])

    # Remaining → line
    ax2 = ax.twinx()
    for i, (sname, sdata) in enumerate(series[1:], start=1):
        vals = [float(sdata.get(l, 0)) for l in labels]
        ax2.plot(x, vals, marker="o", linewidth=2, markersize=4,
                 color=palette[i % len(palette)], label=sname)

    ax.set_xticks(x)
    ax.set_xticklabels(labels)

    lines1, labels1 = ax.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax.legend(lines1 + lines2, labels1 + labels2, fontsize=8, loc="upper left")

    if show_grid:
        ax.yaxis.grid(True, linestyle="--", alpha=0.4)
        ax.set_axisbelow(True)
