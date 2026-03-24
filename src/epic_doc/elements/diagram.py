"""Diagram element — renders graphviz flowcharts as PNG images embedded in DOCX."""
from __future__ import annotations

from typing import TYPE_CHECKING, Any, Dict, List, Optional, Tuple, Union

if TYPE_CHECKING:
    from docx.document import Document
    from epic_doc.styles.theme import Theme
    from epic_doc.utils.tempfiles import TempFileManager

# Node: {"label": str, "shape": str, "color": str | None, "style": str | None}
NodeDef = Dict[str, Any]
# Edge: (from, to) or (from, to, {"label": str, ...})
EdgeDef = Union[Tuple[str, str], Tuple[str, str, Dict[str, Any]]]


def add_flowchart(
    doc: "Document",
    theme: "Theme",
    tmgr: "TempFileManager",
    nodes: Dict[str, Union[str, NodeDef]],
    edges: List[EdgeDef],
    direction: str = "TB",
    width: float = 5.0,
    caption: Optional[str] = None,
    dpi: int = 150,
    graph_attrs: Optional[Dict[str, str]] = None,
) -> None:
    """Render a graphviz flowchart and embed it in the document.

    Args:
        nodes: Dict mapping node_id → label string OR node definition dict.
               Node def keys: label, shape (oval|box|diamond|cylinder|parallelogram),
               color (hex without #), style (filled|dashed|bold).
        edges: List of (from_id, to_id) or (from_id, to_id, attrs_dict).
               attrs_dict keys: label, style (dashed|bold), color.
        direction: Graph layout direction — TB | LR | RL | BT.
        width: Embedded image width in inches.
        caption: Optional caption below the image.
        dpi: Render resolution (default 150).
        graph_attrs: Extra graphviz graph-level attributes.
    """
    try:
        import graphviz
    except ImportError as exc:
        raise ImportError(
            "The 'graphviz' Python package is required for add_flowchart().\n"
            "Install it with: pip install graphviz\n"
            "You also need the Graphviz system binaries:\n"
            "  macOS:  brew install graphviz\n"
            "  Ubuntu: sudo apt install graphviz\n"
            "  Windows: https://graphviz.org/download/"
        ) from exc

    dot = graphviz.Digraph(format="png")
    dot.attr(rankdir=direction, bgcolor="white", dpi=str(dpi))
    dot.attr("graph", fontname="Helvetica", fontsize="11",
             **(graph_attrs or {}))
    dot.attr("node",  fontname="Helvetica", fontsize="10")
    dot.attr("edge",  fontname="Helvetica", fontsize="9")

    # ── Nodes ────────────────────────────────────────────────────────────────
    primary_hex   = theme.primary
    accent_hex    = theme.accent
    header_text   = theme.table_header_text

    _SHAPE_MAP = {
        "oval":         "ellipse",
        "box":          "box",
        "diamond":      "diamond",
        "cylinder":     "cylinder",
        "parallelogram": "parallelogram",
        "ellipse":      "ellipse",
        "rectangle":    "box",
        "rounded":      "box",
    }

    for node_id, node_def in nodes.items():
        if isinstance(node_def, str):
            node_def = {"label": node_def}

        label  = node_def.get("label", node_id)
        shape  = _SHAPE_MAP.get(node_def.get("shape", "box"), "box")
        color  = node_def.get("color")  # custom fill hex
        nstyle = node_def.get("style", "filled")

        # Determine fill color
        if color:
            fill = "#" + color.lstrip("#")
            font_color = "white"
        elif shape == "diamond":
            fill = "#" + accent_hex
            font_color = "white"
        elif shape == "ellipse":
            fill = "#" + primary_hex
            font_color = "white"
        else:
            fill = "#" + theme.table_stripe_bg
            font_color = "#" + theme.body_text

        dot.node(
            node_id,
            label=label,
            shape=shape,
            style=nstyle,
            fillcolor=fill,
            fontcolor=font_color,
            color="#" + theme.table_border,
            penwidth="1.5",
        )

    # ── Edges ────────────────────────────────────────────────────────────────
    for edge in edges:
        if len(edge) == 2:
            src, dst = edge
            attrs: Dict[str, Any] = {}
        else:
            src, dst, attrs = edge[0], edge[1], edge[2]

        edge_label  = attrs.get("label", "")
        edge_style  = attrs.get("style", "solid")
        edge_color  = attrs.get("color", "#" + theme.accent)

        dot.edge(
            src, dst,
            label=edge_label,
            color=edge_color,
            style=edge_style,
            arrowsize="0.8",
            fontsize="9",
        )

    # ── Render ───────────────────────────────────────────────────────────────
    import os
    tmp_path = tmgr.new_png()
    # graphviz renders to <path> (without extension) when we call render
    base_path = tmp_path.replace(".png", "_gv")
    dot.render(filename=base_path, cleanup=True, quiet=True)
    rendered = base_path + ".png"

    # Embed
    from docx.shared import Inches
    doc.add_picture(rendered, width=Inches(width))

    # Track rendered file for cleanup
    tmgr._paths.append(rendered)

    # Caption
    if caption:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor

        def _rgb(h: str) -> RGBColor:
            h = h.lstrip("#")
            return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(caption)
        run.font.italic = True
        run.font.size = Pt(theme.caption_size)
        run.font.color.rgb = _rgb(theme.light_text)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(8)
