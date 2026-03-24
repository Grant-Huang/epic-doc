"""Image element — insert external images with optional caption."""
from __future__ import annotations

import io
from typing import TYPE_CHECKING, Optional, Union

if TYPE_CHECKING:
    from docx.document import Document
    from epic_doc.styles.theme import Theme


def add_image(
    doc: "Document",
    theme: "Theme",
    source: Union[str, bytes],
    width: Optional[float] = None,
    height: Optional[float] = None,
    align: str = "center",
    caption: Optional[str] = None,
) -> None:
    """Insert an image from a file path or bytes.

    Args:
        source: File path string or raw image bytes.
        width: Width in inches. If both width and height are None, uses 4.0".
        height: Height in inches. If only width given, aspect ratio is preserved.
        align: Paragraph alignment — left | center | right.
        caption: Optional italic caption below the image.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    def _rgb(h: str) -> RGBColor:
        h = h.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    _align_map = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }

    if width is None and height is None:
        width = 4.0

    kwargs = {}
    if width is not None:
        kwargs["width"] = Inches(width)
    if height is not None:
        kwargs["height"] = Inches(height)

    if isinstance(source, bytes):
        source = io.BytesIO(source)

    # python-docx adds the picture in a new paragraph
    para = doc.add_paragraph()
    para.alignment = _align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    run = para.add_run()
    run.add_picture(source, **kwargs)

    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.font.italic = True
        cap_run.font.size = Pt(theme.caption_size)
        cap_run.font.color.rgb = _rgb(theme.light_text)
        cap_para.paragraph_format.space_before = Pt(2)
        cap_para.paragraph_format.space_after = Pt(8)
