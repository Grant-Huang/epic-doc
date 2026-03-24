"""Layout elements: page breaks, section breaks, headers, footers, TOC."""
from __future__ import annotations

from typing import TYPE_CHECKING, Optional

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

if TYPE_CHECKING:
    from docx.document import Document
    from epic_doc.styles.theme import Theme


def _rgb(h: str) -> RGBColor:
    h = h.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def add_page_break(doc: "Document") -> None:
    """Insert an explicit page break."""
    doc.add_page_break()


def add_section_break(doc: "Document", break_type: str = "next_page") -> None:
    """Insert a section break.

    break_type: next_page | even_page | odd_page | continuous
    """
    from docx.enum.section import WD_SECTION
    _map = {
        "next_page":  WD_SECTION.NEW_PAGE,
        "even_page":  WD_SECTION.EVEN_PAGE,
        "odd_page":   WD_SECTION.ODD_PAGE,
        "continuous": WD_SECTION.CONTINUOUS,
    }
    doc.add_section(_map.get(break_type, WD_SECTION.NEW_PAGE))


def set_header(
    doc: "Document",
    theme: "Theme",
    text: str,
    align: str = "right",
    include_title: bool = False,
) -> None:
    """Set document header text for the default section."""
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    header = section.header
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    _align_map = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = _align_map.get(align, WD_ALIGN_PARAGRAPH.RIGHT)

    for run in para.runs:
        run.text = ""

    run = para.add_run(text)
    run.font.name = theme.body_font
    run.font.size = Pt(theme.caption_size)
    run.font.color.rgb = _rgb(theme.light_text)

    # Thin bottom border on header
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "2")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), theme.table_grid)
    pBdr.append(bottom)
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)


def set_footer(
    doc: "Document",
    theme: "Theme",
    text: Optional[str] = None,
    page_number: bool = True,
    align: str = "center",
) -> None:
    """Set document footer with optional page number field."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    _align_map = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = _align_map.get(align, WD_ALIGN_PARAGRAPH.CENTER)

    for run in para.runs:
        run.text = ""

    def _styled_run(t: str) -> None:
        r = para.add_run(t)
        r.font.name = theme.body_font
        r.font.size = Pt(theme.caption_size)
        r.font.color.rgb = _rgb(theme.light_text)

    if text:
        _styled_run(text)
        if page_number:
            _styled_run("   |   ")

    if page_number:
        _styled_run("第 ")
        _insert_page_number(para, theme)
        _styled_run(" 页")

    # Top border on footer
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "2")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), theme.table_grid)
    pBdr.append(top)
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pPr.append(pBdr)


def _insert_page_number(paragraph, theme: "Theme") -> None:
    """Insert a PAGE field run into a paragraph."""
    run = paragraph.add_run()
    run.font.name = theme.body_font
    run.font.size = Pt(theme.caption_size)
    run.font.color.rgb = _rgb(theme.light_text)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    instrText.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_toc(
    doc: "Document",
    theme: "Theme",
    title: str = "目录",
    depth: int = 3,
) -> None:
    """Insert a TOC placeholder (requires Word / LibreOffice to refresh with F9)."""
    from epic_doc.utils.xml_helpers import add_toc_field

    # Title paragraph
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run(title)
    run.font.name = theme.heading_font
    run.font.size = Pt(theme.h2_size)
    run.font.bold = True
    run.font.color.rgb = _rgb(theme.primary)
    title_para.paragraph_format.space_before = Pt(theme.h2_space_before)
    title_para.paragraph_format.space_after = Pt(theme.h2_space_after)

    # TOC field
    add_toc_field(doc)
    doc.add_paragraph()  # spacer


def set_metadata(
    doc: "Document",
    title: Optional[str] = None,
    author: Optional[str] = None,
    subject: Optional[str] = None,
    description: Optional[str] = None,
) -> None:
    """Set core document metadata properties."""
    props = doc.core_properties
    if title:
        props.title = title
    if author:
        props.author = author
    if subject:
        props.subject = subject
    if description:
        props.description = description
