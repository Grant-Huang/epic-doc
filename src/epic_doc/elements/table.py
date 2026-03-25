"""Complex table element with merging, shading, borders, and preset styles."""
from __future__ import annotations

from typing import TYPE_CHECKING, Any, List, Optional, Tuple

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from epic_doc.utils.xml_helpers import (
    set_cell_bg,
    set_cell_borders,
    set_cell_vertical_alignment,
    set_table_borders,
    set_table_no_borders,
)

if TYPE_CHECKING:
    from docx.document import Document

    from epic_doc.styles.theme import Theme

# Table style presets
_STYLES = ("striped", "grid", "minimal", "bordered", "dark", "card")

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _cell_text(cell, text: str, theme: "Theme", font_size: int,
               bold: bool = False, color: Optional[str] = None,
               align: str = "left") -> None:
    for para in cell.paragraphs:
        para._element.getparent().remove(para._element)
    para = cell.add_paragraph()
    para.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(str(text))
    run.font.name = theme.body_font
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = _rgb(color)


def add_table(
    doc: "Document",
    theme: "Theme",
    data: List[List[Any]],
    headers: bool = True,
    style: str = "striped",
    col_widths: Optional[List[float]] = None,
    merge: Optional[List[Tuple[int, int, int, int]]] = None,
    caption: Optional[str] = None,
    align: str = "left",
    font_size: Optional[int] = None,
    cell_align: str = "left",
) -> None:
    """Add a feature-rich table to the document.

    Args:
        data: 2-D list of cell values (row-major). First row is header if ``headers=True``.
        headers: Whether the first row should be styled as a header.
        style: One of ``striped`` | ``grid`` | ``minimal`` | ``bordered`` | ``dark``.
        col_widths: Column widths in inches. None → equal distribution across 6".
        merge: List of (start_row, start_col, end_row, end_col) merge regions.
        caption: Optional caption paragraph below the table.
        align: Table horizontal alignment: ``left`` | ``center`` | ``right``.
        font_size: Override body font size for all cells.
        cell_align: Default text alignment inside cells.
    """
    if not data:
        return

    fs = font_size or theme.body_size
    n_rows = len(data)
    n_cols = max(len(row) for row in data)

    # Pad rows that are shorter than n_cols
    padded = [list(row) + [""] * (n_cols - len(row)) for row in data]

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    # --- Column widths ---
    total_width = 6.0  # inches, default content area
    if col_widths:
        widths = col_widths[:n_cols]
        widths += [1.0] * (n_cols - len(widths))
    else:
        w = total_width / n_cols
        widths = [w] * n_cols

    for col_idx, width in enumerate(widths):
        for row in table.rows:
            row.cells[col_idx].width = Inches(width)

    # --- Table alignment ---
    from docx.enum.table import WD_TABLE_ALIGNMENT
    align_map = {
        "left": WD_TABLE_ALIGNMENT.LEFT,
        "center": WD_TABLE_ALIGNMENT.CENTER,
        "right": WD_TABLE_ALIGNMENT.RIGHT,
    }
    table.alignment = align_map.get(align, WD_TABLE_ALIGNMENT.LEFT)

    # --- Fill cells ---
    for r_idx, row_data in enumerate(padded):
        for c_idx, value in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            is_header_row = headers and r_idx == 0

            if is_header_row:
                _cell_text(cell, value, theme, fs, bold=True,
                           color=theme.table_header_text, align="center")
                set_cell_bg(cell, theme.table_header_bg)
            else:
                _cell_text(cell, value, theme, fs, align=cell_align)
                # Striped rows
                if style == "striped" and r_idx % 2 == 0:
                    set_cell_bg(cell, theme.table_stripe_bg)

            set_cell_vertical_alignment(cell, "center")

    # --- Apply style-specific borders ---
    if style in ("grid", "bordered"):
        set_table_borders(table, color=theme.table_border, size="4")
    elif style == "minimal":
        set_table_no_borders(table)
        # Only draw top/bottom for each row
        for r_idx in range(n_rows):
            for c_idx in range(n_cols):
                cell = table.cell(r_idx, c_idx)
                set_cell_borders(cell, top=True, bottom=True, left=False, right=False,
                                 color=theme.table_grid, size="2")
    elif style == "dark":
        set_table_borders(table, color=theme.table_header_bg, size="6")
        for r_idx in range(1, n_rows):
            for c_idx in range(n_cols):
                cell = table.cell(r_idx, c_idx)
                if r_idx % 2 == 0:
                    set_cell_bg(cell, theme.table_stripe_bg)
    elif style == "card":
        # Card-like block: strong outer border, consistent light body background.
        set_table_borders(table, color=theme.table_border, size="6")
        for r_idx in range(n_rows):
            for c_idx in range(n_cols):
                cell = table.cell(r_idx, c_idx)
                # Keep header row dark when headers=True; shade body rows uniformly.
                if not (headers and r_idx == 0):
                    set_cell_bg(cell, theme.table_stripe_bg)
    elif style == "striped":
        set_table_borders(table, color=theme.table_border, size="4")

    # --- Merge cells ---
    if merge:
        for sr, sc, er, ec in merge:
            a = table.cell(sr, sc)
            b = table.cell(er, ec)
            a.merge(b)

    # --- Caption ---
    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run(caption)
        cap_run.font.italic = True
        cap_run.font.size = Pt(theme.caption_size)
        cap_run.font.color.rgb = _rgb(theme.light_text)
        cap_para.paragraph_format.space_before = Pt(2)
        cap_para.paragraph_format.space_after = Pt(8)
