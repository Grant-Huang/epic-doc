"""Text elements: headings, paragraphs, lists, hyperlinks, code blocks, callouts."""
from __future__ import annotations

from typing import TYPE_CHECKING, List, Optional, Union

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from epic_doc.utils.xml_helpers import add_paragraph_border_bottom, make_hyperlink

if TYPE_CHECKING:
    from docx.document import Document

    from epic_doc.styles.theme import Theme

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def add_heading(
    doc: "Document",
    theme: "Theme",
    text: str,
    level: int = 1,
    align: str = "left",
) -> None:
    """Add a styled heading to the document."""
    level = max(1, min(4, level))
    para = doc.add_heading(text, level=level)
    para.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    run = para.runs[0] if para.runs else None

    color_map = {
        1: theme.primary,
        2: theme.secondary,
        3: theme.accent,
        4: theme.accent,
    }
    size_map = {
        1: theme.h1_size,
        2: theme.h2_size,
        3: theme.h3_size,
        4: theme.h4_size,
    }
    bold_map = {1: theme.h1_bold, 2: theme.h2_bold, 3: theme.h3_bold, 4: theme.h4_bold}
    italic_map = {1: theme.h1_italic, 2: theme.h2_italic, 3: theme.h3_italic, 4: False}

    if run:
        run.font.color.rgb = _rgb(color_map[level])
        run.font.size = Pt(size_map[level])
        run.font.name = theme.heading_font
        run.font.bold = bold_map[level]
        run.font.italic = italic_map[level]
        if level == 1 and theme.h1_caps:
            run.font.all_caps = True

    # Spacing
    pf = para.paragraph_format
    space_before = {1: theme.h1_space_before, 2: theme.h2_space_before,
                    3: theme.h3_space_before, 4: theme.h4_space_before}
    space_after  = {1: theme.h1_space_after,  2: theme.h2_space_after,
                    3: theme.h3_space_after,  4: theme.h4_space_after}
    pf.space_before = Pt(space_before[level])
    pf.space_after  = Pt(space_after[level])

    # H1 decorative bottom border
    if level == 1 and theme.h1_border:
        add_paragraph_border_bottom(para, theme.h1_border_hex, size="6")


def add_paragraph(
    doc: "Document",
    theme: "Theme",
    text: str,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    color: Optional[str] = None,
    align: str = "left",
    font_size: Optional[int] = None,
    style: Optional[str] = None,
) -> None:
    """Add a body paragraph with optional inline formatting."""
    para = doc.add_paragraph(style=style or "Normal")
    para.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    pf = para.paragraph_format
    pf.space_before = Pt(theme.body_space_before)
    pf.space_after  = Pt(theme.body_space_after)

    run = para.add_run(text)
    run.font.name = theme.body_font
    run.font.size = Pt(font_size or theme.body_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    run.font.color.rgb = _rgb(color.lstrip("#") if color else theme.body_text)


def add_list(
    doc: "Document",
    theme: "Theme",
    items: List[Union[str, list]],
    style: str = "bullet",
    level: int = 0,
) -> None:
    """Add a list (bullet or numbered), supporting nested items.

    Nested items are represented as a sub-list (Python list) within the items list.
    """
    docx_style = "List Bullet" if style == "bullet" else "List Number"

    for item in items:
        if isinstance(item, list):
            # Recurse for nested list at level+1
            add_list(doc, theme, item, style=style, level=level + 1)
        else:
            para = doc.add_paragraph(style=docx_style)
            run = para.add_run(str(item))
            run.font.name = theme.body_font
            run.font.size = Pt(theme.body_size)
            run.font.color.rgb = _rgb(theme.body_text)
            # Indent nested levels
            if level > 0:
                para.paragraph_format.left_indent = Pt(18 * level)


def add_hyperlink(
    doc: "Document",
    theme: "Theme",
    text: str,
    url: str,
) -> None:
    """Add a paragraph containing a hyperlink."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(theme.body_space_before)
    pf.space_after  = Pt(theme.body_space_after)
    make_hyperlink(para, text, url, color=theme.accent)


def add_code_block(
    doc: "Document",
    theme: "Theme",
    code: str,
    language: Optional[str] = None,
) -> None:
    """Add a styled monospace code block inside a shaded box."""
    from epic_doc.utils.xml_helpers import set_cell_bg

    # Use a 1×1 table as the code container for background shading
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)

    # Background color
    set_cell_bg(cell, theme.code_bg)

    # Clear default empty paragraph and add code lines
    for para in cell.paragraphs:
        p = para._element
        p.getparent().remove(p)

    for i, line in enumerate(code.splitlines()):
        para = cell.add_paragraph()
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after  = Pt(0)
        run = para.add_run(line if line else " ")
        run.font.name = theme.mono_font
        run.font.size = Pt(theme.code_size)
        run.font.color.rgb = _rgb(theme.code_text)

    # Add spacing paragraph after table
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_callout(
    doc: "Document",
    theme: "Theme",
    text: str,
    style: str = "info",
    title: Optional[str] = None,
) -> None:
    """Add a highlighted callout box (info | warning | danger | success)."""
    from epic_doc.utils.xml_helpers import set_cell_bg, set_cell_borders

    color_map = {
        "info":    (theme.callout.info_bg,    theme.callout.info_border),
        "warning": (theme.callout.warning_bg, theme.callout.warning_border),
        "danger":  (theme.callout.danger_bg,  theme.callout.danger_border),
        "success": (theme.callout.success_bg, theme.callout.success_border),
    }
    bg_color, border_color = color_map.get(style, color_map["info"])

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_color)
    set_cell_borders(cell, left=True, top=False, bottom=False, right=False,
                     color=border_color, size="12")

    for para in cell.paragraphs:
        para._element.getparent().remove(para._element)

    if title:
        title_para = cell.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.bold = True
        title_run.font.size = Pt(theme.body_size)
        title_run.font.name = theme.body_font
        title_run.font.color.rgb = _rgb(border_color)
        title_para.paragraph_format.space_after = Pt(2)

    body_para = cell.add_paragraph()
    body_run = body_para.add_run(text)
    body_run.font.size = Pt(theme.body_size)
    body_run.font.name = theme.body_font
    body_run.font.color.rgb = _rgb(theme.body_text)
    body_para.paragraph_format.space_after = Pt(0)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_horizontal_rule(doc: "Document", theme: "Theme") -> None:
    """Add a thin horizontal rule paragraph."""
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(6)
    add_paragraph_border_bottom(para, theme.accent, size="4")
